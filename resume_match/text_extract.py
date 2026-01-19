"""
Text extraction module for resumes in various formats.
Supports: TXT, DOCX, PDF

Architecture:
- extract_text() - Unified interface, auto-detects file type
- _read_txt()    - Plain text extraction
- _read_docx()   - DOCX extraction with table support
- _read_pdf()    - PDF extraction with artifact cleaning
- _postprocess_text() - Common text normalization
"""

import os
import re
from typing import Tuple


def _postprocess_text(text: str) -> str:
    """
    Common text normalization applied to all extracted text.

    - Normalize whitespace
    - Remove empty lines
    - Ensure consistent line breaks
    """
    # Normalize line breaks
    text = text.replace('\r\n', '\n').replace('\r', '\n')

    # Remove excessive whitespace
    lines = []
    for line in text.split('\n'):
        line = line.strip()
        if line:
            # Normalize internal whitespace
            line = re.sub(r'\s+', ' ', line)
            lines.append(line)

    return '\n'.join(lines)


def _read_txt(path: str) -> str:
    """Extract text from plain text file."""
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()
    return _postprocess_text(text)


def _read_docx(path: str) -> str:
    """
    Extract text from DOCX file.
    Handles both paragraphs and tables.
    """
    from docx import Document  # type: ignore

    doc = Document(path)
    parts = []

    # Extract paragraphs
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)

    # Extract tables
    # Join cells with space (not pipe) to avoid splitting issues
    for table in doc.tables:
        for row in table.rows:
            cells = []
            for cell in row.cells:
                ct = (cell.text or "").strip()
                if ct:
                    cells.append(ct)
            if cells:
                # Merge cells with space separator
                parts.append(" ".join(cells))

    full_text = "\n".join(parts)
    return _postprocess_text(full_text)


def _clean_pdf_artifacts(text: str) -> str:
    """
    Clean up common PDF extraction artifacts:
    1. Remove hyphenation at line breaks (tech-\nnical -> technical)
    2. Intelligently merge continuation lines
    3. Preserve paragraph boundaries
    4. Normalize whitespace
    """
    # Step 1: Handle hyphenation at line breaks
    # Pattern: word- followed by newline and then word continuation
    text = re.sub(r'(\w)-\s*\n\s*(\w)', r'\1\2', text)

    # Step 2: Smart line merging
    lines = text.split('\n')
    merged_lines = []
    i = 0

    while i < len(lines):
        current_line = lines[i].strip()

        # Skip empty lines (preserve as paragraph breaks)
        if not current_line:
            merged_lines.append('')
            i += 1
            continue

        # Check if this line should be merged with next line
        should_merge = False
        if i + 1 < len(lines):
            next_line = lines[i + 1].strip()

            # Don't merge if current line ends with strong punctuation
            if current_line and current_line[-1] in '.!?:':
                should_merge = False
            # Don't merge if next line is empty (paragraph break)
            elif not next_line:
                should_merge = False
            # Don't merge if next line looks like a header
            elif next_line.isupper() and len(next_line) < 50:
                should_merge = False
            # Merge if next line starts with lowercase (likely continuation)
            elif next_line and next_line[0].islower():
                should_merge = True
            # Merge if current line is short and looks incomplete
            elif len(current_line) < 80 and current_line[-1] not in '.!?':
                should_merge = True

        if should_merge:
            next_line = lines[i + 1].strip()
            merged_lines.append(current_line + ' ' + next_line)
            i += 2  # Skip next line since we merged it
        else:
            merged_lines.append(current_line)
            i += 1

    # Rejoin with newlines
    text = '\n'.join(merged_lines)

    # Step 3: Normalize multiple newlines
    text = re.sub(r'\n{3,}', '\n\n', text)

    return text.strip()


def _read_pdf(path: str) -> str:
    """
    Extract text from PDF file.
    Handles encrypted PDFs and scanned documents.
    """
    from pypdf import PdfReader  # type: ignore

    try:
        reader = PdfReader(path)

        # Check if PDF is encrypted
        if reader.is_encrypted:
            raise ValueError(
                "PDF is password-protected. "
                "Please remove password protection or use an unprotected version."
            )

        # Extract text from all pages
        parts = []
        for page in reader.pages:
            text = page.extract_text() or ""
            text = text.strip()
            if text:
                parts.append(text)

        # Combine all pages
        full_text = "\n\n".join(parts)

        # Clean up PDF artifacts
        full_text = _clean_pdf_artifacts(full_text)

        # Apply common postprocessing
        full_text = _postprocess_text(full_text)

        # Heuristic: if very short, likely scanned/no text layer
        if len(full_text.strip()) < 200:
            raise ValueError(
                "PDF text extraction returned very little text. "
                "This PDF may be:\n"
                "  1. Scanned (image-only, requires OCR)\n"
                "  2. Has complex layout that failed to extract\n"
                "  3. Encrypted with text extraction restrictions\n"
                "Please try:\n"
                "  - Converting to DOCX format\n"
                "  - Saving as text-based PDF\n"
                "  - Using OCR if it's a scanned document"
            )

        return full_text

    except Exception as e:
        error_msg = str(e).lower()
        if "encrypted" in error_msg or "password" in error_msg:
            raise ValueError(f"PDF encryption error: {str(e)}")
        else:
            raise ValueError(
                f"PDF extraction failed: {str(e)}\n"
                "Try converting the PDF to DOCX or TXT format."
            )


def extract_text(path: str) -> Tuple[str, str]:
    """
    Unified text extraction interface.
    Auto-detects file type and applies appropriate extraction method.

    Args:
        path: Path to the resume file

    Returns:
        Tuple of (extracted_text, extraction_method)
        extraction_method is one of: "txt", "docx", "pdf"

    Raises:
        FileNotFoundError: If file doesn't exist
        ValueError: If file type is unsupported or extraction fails

    Example:
        text, method = extract_text("resume.pdf")
        print(f"Extracted via {method}: {len(text)} characters")
    """
    if not os.path.exists(path):
        raise FileNotFoundError(f"File not found: {path}")

    ext = os.path.splitext(path.lower())[1]

    if ext == ".txt":
        return _read_txt(path), "txt"
    elif ext == ".docx":
        return _read_docx(path), "docx"
    elif ext == ".pdf":
        return _read_pdf(path), "pdf"
    else:
        raise ValueError(
            f"Unsupported file type: {ext}\n"
            f"Supported formats: .txt, .docx, .pdf"
        )